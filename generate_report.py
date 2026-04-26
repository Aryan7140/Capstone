from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('FraudShield: Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('AI-Powered Real-Time Digital Arrest & Fraud Detection System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')

    # --- Section 1: Problem Statement ---
    doc.add_heading('1. Problem Statement', level=1)
    doc.add_paragraph(
        'In recent years, India has seen a massive surge in "Digital Arrest" and cyber-impersonation scams. '
        'Fraudsters impersonate high-ranking officials from agencies like the CBI, Narcotics Department, '
        'and State Police to intimidate citizens into transferring large sums of money. '
        'The core problem lies in the victim\'s inability to verify the authenticity of a threat in real-time, '
        'leading to multi-crore financial losses and psychological trauma.'
    )

    # --- Section 2: Proposed Solution ---
    doc.add_heading('2. The FraudShield Solution', level=1)
    doc.add_paragraph(
        'FraudShield is an integrated AI-driven ecosystem designed to dismantle the digital arrest scam lifecycle. '
        'It provides a three-layered defense strategy:'
    )
    doc.add_paragraph('• Real-time NLP Analysis: Detects intimidation patterns and impersonation keywords in both English and romanized Hindi (Hinglish).', style='List Bullet')
    doc.add_paragraph('• Transactional Risk Profiling: Instantly cross-references mentioned bank accounts against a historical database of 20,000+ fraud-linked transaction records.', style='List Bullet')
    doc.add_paragraph('• Centralized Threat Database: Every detected scam is logged to a global Firebase database, preventing scammers from using the same tactics or account numbers across different victims.', style='List Bullet')

    # --- Section 3: Technical Implementation ---
    doc.add_heading('3. Technical Implementation', level=1)
    
    doc.add_heading('3.1 Full-Stack Architecture', level=2)
    doc.add_paragraph('• Frontend: A responsive React.js Dashboard providing real-time visual alerts and investigation tools.', style='List Bullet')
    doc.add_paragraph('• Backend: A high-performance Flask API utilizing Python’s data science stack (Pandas, Scikit-Learn, XGBoost).', style='List Bullet')
    doc.add_paragraph('• Database Layer: Firebase Real-time Database for immediate persistence and global threat syncing.', style='List Bullet')

    doc.add_heading('3.2 Data Optimization Engine', level=2)
    doc.add_paragraph(
        'To overcome cloud hosting limitations, the system uses "Parquet" binary storage and background loading threads. '
        'This allows the application to handle large datasets while starting up in under 30 seconds, ensuring high availability.'
    )

    # --- Section 4: Attainment of Objectives ---
    doc.add_heading('4. Attainment of Defined Objectives', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project Objective'
    hdr_cells[1].text = 'Status / Attainment'
    
    objs = [
        ('Real-time Scam Detection', 'Achieved. The hybrid ML engine flags Digital Arrest patterns with 90%+ confidence.'),
        ('Multilingual Support', 'Achieved. The system supports English and Hinglish (e.g., "ek lakh", "bol raha hoon").'),
        ('Financial Forensics', 'Achieved. Integrated lookup for 20k+ records with risk-score generation.'),
        ('Scalable Deployment', 'Achieved. Fully operational on Render/Vercel with optimized memory mapping.')
    ]
    
    for obj, status in objs:
        row_cells = table.add_row().cells
        row_cells[0].text = obj
        row_cells[1].text = status

    # --- Section 5: Dashboard and Alerts Usage ---
    doc.add_heading('5. Dashboard and Alert System Usage', level=1)
    doc.add_paragraph(
        'The FraudShield Dashboard serves as the central "Command Center" for the user, '
        'providing a real-time window into the security health of the ecosystem.'
    )
    
    doc.add_heading('5.1 The Unified Dashboard', level=2)
    doc.add_paragraph(
        'The dashboard provides high-level metrics through dynamic visualizations. '
        'Users can instantly see the total count of identified fraud cases, suspicious alerts, '
        'and legitimate transactions. This bird’s-eye view is critical for security analysts '
        'to identify spikes in fraudulent activity.'
    )
    
    doc.add_heading('5.2 Real-Time Alerts Feed', level=2)
    doc.add_paragraph(
        'The Alerts system provides a live, granular feed of every flagged attempt. '
        'Each alert includes the transaction ID, the risk score, and the specific reason for '
        'the flag (e.g., "Impersonation detected" or "Fraud-linked account"). '
        'This allows the user to take immediate action, such as freezing an account or '
        'notifying authorities.'
    )
    
    doc.add_heading('5.3 Interactive Investigation', level=2)
    doc.add_paragraph(
        'Beyond passive monitoring, the user can use the "Investigate" module to manually '
        'verify any suspicious communication. By entering a message, the dashboard provides '
        'an instant "Threat Level" assessment, protecting the user from making risky decisions '
        'under pressure.'
    )

    # Save the document
    doc.save('FraudShield_Technical_Report.docx')
    print("Report generated successfully: FraudShield_Technical_Report.docx")

if __name__ == "__main__":
    create_report()
